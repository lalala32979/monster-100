import openpyxl
import shutil

from docx import Document
from docx.shared import Inches
from docx.enum.shape import WD_INLINE_SHAPE_TYPE
from docx.shared import Cm, Inches, Pt

# 打开 Excel 文件
workbook = openpyxl.load_workbook('moster_data.xlsx')
sheet = workbook.active

type_lis = ['小怪','精英','BOSS']

count = 0
# 遍历 Key 列和 C 列
for cell in sheet['A']:
    # if isinstance(cell.value, (int,float)):
        
    if cell.value is not None and cell.row > 4:
        # 获取 C 列的数据
        c_cell = cell.offset(column=2)
        folder_name = c_cell.value
        d_cell = cell.offset(column=3)
        folder_name = c_cell.value
        folder_type = type_lis[d_cell.value - 1]
        
        num_str = str(cell.value).zfill(3)
        
        word_file = '模板2.docx'
        # 复制 Word 文件
        shutil.copy(word_file, f'./world/{num_str}{folder_name}.docx')

        # 打开 Word 文件
        document = Document(f'world/{num_str}{folder_name}.docx')

        # 将文本替换为表格中的值
        for paragraph in document.paragraphs:
            if 'name' in paragraph.text:
                paragraph.text = paragraph.text.replace('name', folder_name)
            if 'type' in paragraph.text:
                paragraph.text = paragraph.text.replace('type', folder_type)

        new_image_path = f'抠图128X128/{folder_name}.png'
        document.add_picture(new_image_path, width=Inches(2.0), height=Inches(2.0))

        document.add_paragraph(f'近战{folder_type}，拥有较大的近战警戒/取消警戒范围(800/1400)')

        # 遍历所有 InlineShape 对象
        # for image in document.inline_shapes:
        #     if image._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name == '图片 2':
        #         new_image_path = f'抠图128X128/{folder_name}.png'
        #         image._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name = '图片 2'
        #         image._inline.graphic.graphicData.pic.nvPicPr.cNvPr.descr = folder_name
        #         image._inline.graphic.graphicData.pic.blipFill.blip.embed = document.add_picture(new_image_path)._inline.graphic.graphicData.pic.blipFill.blip.embed
        #         image.width = Pt(128)
        #         image.height = Pt(128)
        #         break
            
            # if image._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name == '图片 2':
            #     new_image_path = '不死将军.png'
            #     # 删除原有的图片
            #     # image._inline.remove()
            #     # 插入新的图片
            #     paragraph = document.add_paragraph()
            #     paragraph.add_run().add_picture(new_image_path, width=Pt(128), height=Pt(128))  # 插入图片并设置宽度和高度
            #     break

        # 保存 Word 文件
        document.save(f'world/{num_str}{folder_name}.docx')
        
        count += 1
        if count == 10:
            break